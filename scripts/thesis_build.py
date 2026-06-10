"""Word document builders for Corner Store graduation thesis."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


AUTHOR = "[Your Full Name]"
SUPERVISOR = "Dr. Hala Abd Elgelil"
UNIVERSITY = "Helwan University"
FACULTY = "[Faculty of Computers and Artificial Intelligence]"
DEPARTMENT = "[Department of Computer Science]"
DEGREE = "Bachelor of Science"
DATE = "June 2026"
ACADEMIC_YEAR = "2025/2026"

TITLE = (
    "Corner Store: An AI-Augmented Full-Stack E-Commerce Platform with "
    "Retrieval-Augmented Conversational Assistance and Multimodal Visual Product Discovery"
)

ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DOCX = ROOT / "CornerStore-Thesis-Helwan-University.docx"
OUTPUT_PDF = ROOT / "GRADUATION_PROJECT_DOCUMENTATION.pdf"


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(6)

    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        if level == 1:
            style.font.size = Pt(16)
        elif level == 2:
            style.font.size = Pt(14)
        else:
            style.font.size = Pt(13)

    return doc


def add_centered(doc: Document, text: str, *, bold: bool = False, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_figure(doc: Document, caption: str, diagram: str) -> None:
    """Legacy text figure — prefer add_figure_image for visual diagrams."""
    add_paragraph(doc, caption)
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)


def add_figure_image(
    doc: Document,
    caption: str,
    image_path: Path | str,
    *,
    width_inches: float = 6.2,
) -> None:
    """Embed a rendered diagram PNG with caption."""
    add_paragraph(doc, caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)


def add_screenshot_placeholder(doc: Document, caption: str, description: str) -> None:
    add_paragraph(doc, caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ INSERT SCREENSHOT: {description} ]")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    for _ in range(3):
        doc.add_paragraph()
    add_paragraph(
        doc,
        f"Insert screenshot: {description.lower()}.",
    )


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    add_centered(doc, UNIVERSITY, bold=True, size=14)
    add_centered(doc, FACULTY, size=12)
    add_centered(doc, DEPARTMENT, size=12)
    doc.add_paragraph()
    add_centered(doc, TITLE, bold=True, size=13)
    doc.add_paragraph()
    add_centered(doc, "A Graduation Project Thesis", size=12)
    add_centered(
        doc,
        "Submitted in Partial Fulfillment of the Requirements for the Degree of",
        size=12,
    )
    add_centered(doc, DEGREE, bold=True, size=12)
    doc.add_paragraph()
    add_centered(doc, f"By: {AUTHOR}", size=12)
    add_centered(doc, f"Supervisor: {SUPERVISOR}", size=12)
    doc.add_paragraph()
    add_centered(doc, f"Academic Year: {ACADEMIC_YEAR}", size=12)
    add_centered(doc, DATE, size=12)
    add_page_break(doc)
