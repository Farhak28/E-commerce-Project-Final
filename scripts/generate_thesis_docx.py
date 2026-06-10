#!/usr/bin/env python3
"""Generate Corner Store graduation thesis — verified content only."""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
ROOT = SCRIPT_DIR.parent
sys.path.insert(0, str(SCRIPT_DIR))

from thesis_build import add_cover_page, setup_document
from thesis_content import build_all_sections
from thesis_diagram_render import render_all

OUTPUT_DOCX = ROOT / "CornerStore-Thesis-Helwan-University.docx"
OUTPUT_PDF = ROOT / "GRADUATION_PROJECT_DOCUMENTATION.pdf"
FALLBACK_DOCX = ROOT / "CornerStore-Thesis-Helwan-University-NEW.docx"


def generate_docx(path: Path) -> Path:
    print("Rendering UML diagrams (matplotlib)...")
    diagrams = render_all()
    print(f"  {len(diagrams)} diagram PNGs -> scripts/diagram_output/")

    doc = setup_document()
    add_cover_page(doc)
    build_all_sections(doc)
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(path))
        return path
    except PermissionError:
        doc.save(str(FALLBACK_DOCX))
        print(f"WARNING: Could not overwrite {path.name} (close Word if open).")
        print(f"         Saved to: {FALLBACK_DOCX}")
        return FALLBACK_DOCX


def word_count(path: Path) -> int:
    from docx import Document

    doc = Document(str(path))
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                words += len(cell.text.split())
    return words


def try_pdf_export(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception:
        pass

    for candidate in (
        "soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
    ):
        try:
            subprocess.run([candidate, "--version"], capture_output=True, check=True, timeout=15)
            subprocess.run(
                [candidate, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
                check=True,
                timeout=120,
                capture_output=True,
            )
            generated = docx_path.with_suffix(".pdf")
            if generated.exists() and generated != pdf_path:
                generated.replace(pdf_path)
            return pdf_path.exists()
        except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
            continue
    return False


if __name__ == "__main__":
    saved = generate_docx(OUTPUT_DOCX)
    wc = word_count(saved)
    print(f"Created: {saved}")
    print(f"Words: {wc:,} (screenshot placeholders add extra pages when filled)")

    pdf_source = saved
    if try_pdf_export(pdf_source, OUTPUT_PDF):
        print(f"Created: {OUTPUT_PDF}")
        try:
            from pypdf import PdfReader
            print(f"PDF pages (excl. screenshots you add): {len(PdfReader(str(OUTPUT_PDF)).pages)}")
        except Exception:
            pass
    else:
        print(f"Open {OUTPUT_DOCX.name} in Word and export PDF to update {OUTPUT_PDF.name}")
